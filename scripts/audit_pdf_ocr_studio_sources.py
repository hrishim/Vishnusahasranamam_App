from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
STUDIO_ROOT = Path("/Users/sathyavasu/Projects/codex/PDF_OCR_Studio/final_outputs")
OUT_DIR = ROOT / "outputs" / "pdf_ocr_studio_sources"

VOLUMES = [
    ("Volume 01", STUDIO_ROOT / "2bb35ea90639" / "VS_Namavali_vol 01Ser_corrected.docx"),
    ("Volume 02", STUDIO_ROOT / "0bcd4185e03a" / "VS_Namavali_vol 02ser_corrected.docx"),
    ("Volume 03", STUDIO_ROOT / "80ad95d0a40b" / "VS_Namavali_vol 03ser_corrected.docx"),
    ("Volume 04", STUDIO_ROOT / "b7adf92b19a2" / "VS_Namavali_vol 04Ser_corrected.docx"),
    ("Volume 05", STUDIO_ROOT / "62bc4bdc15d0" / "VS_Namavali_vol 05ser_corrected.docx"),
]

BAD_CHARS = set("�□■●◆◇�")
SUSPECT_PATTERNS = [
    ("replacement_char", re.compile(r"\ufffd")),
    ("box_char", re.compile(r"[□■]")),
    ("long_symbol_run", re.compile(r"[^\w\s\u0900-\u097F.,;:'\"()|/\\-]{4,}")),
    ("very_long_token", re.compile(r"\S{45,}")),
    ("ocr_noise_token", re.compile(r"\b[A-Za-z]*[0-9][A-Za-z]*[0-9][A-Za-z0-9]*\b")),
    ("broken_common_word", re.compile(r"\b(?:tne|tne|tnat|tneir|tnere|wnen|wnich|wnere)\b", re.I)),
]


@dataclass
class VolumeAudit:
    label: str
    path: Path
    paragraphs: int
    nonempty: int
    chars: int
    devanagari_chars: int
    latin_chars: int
    suspect_lines: list[tuple[int, str, str]]
    samples: list[str]


def text_runs(path: Path) -> list[str]:
    doc = Document(str(path))
    return [p.text.strip() for p in doc.paragraphs]


def audit_volume(label: str, path: Path) -> VolumeAudit:
    lines = text_runs(path)
    nonempty_lines = [line for line in lines if line]
    text = "\n".join(nonempty_lines)
    suspect: list[tuple[int, str, str]] = []
    for idx, line in enumerate(nonempty_lines, start=1):
        reasons = []
        if any(ch in BAD_CHARS for ch in line):
            reasons.append("bad glyph")
        for name, pattern in SUSPECT_PATTERNS:
            if pattern.search(line):
                reasons.append(name)
        ascii_letters = sum(1 for ch in line if "A" <= ch <= "Z" or "a" <= ch <= "z")
        ascii_noise = sum(1 for ch in line if ch in "{}[]<>~`_^=")
        if ascii_noise >= 3 and ascii_noise / max(1, len(line)) > 0.04:
            reasons.append("symbol noise")
        if reasons:
            suspect.append((idx, ", ".join(sorted(set(reasons))), line[:220]))

    return VolumeAudit(
        label=label,
        path=path,
        paragraphs=len(lines),
        nonempty=len(nonempty_lines),
        chars=len(text),
        devanagari_chars=sum(1 for ch in text if "\u0900" <= ch <= "\u097F"),
        latin_chars=sum(1 for ch in text if "A" <= ch <= "Z" or "a" <= ch <= "z"),
        suspect_lines=suspect,
        samples=nonempty_lines[:8],
    )


def add_paragraph_copy(target: Document, text: str) -> None:
    para = target.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.name = "Arial Unicode MS"
    run.font.size = Pt(11)


def make_combined_docx() -> Path:
    combined = Document()
    styles = combined.styles
    styles["Normal"].font.name = "Arial Unicode MS"
    styles["Normal"].font.size = Pt(11)
    combined.add_heading("Vishnusahasranamam - Corrected Scan Text Source", level=0)
    combined.add_paragraph(
        "Combined from PDF_OCR_Studio corrected outputs for scanned volumes 01-05. "
        "This is a source candidate for search/indexing, not an edited publication master."
    )
    for label, path in VOLUMES:
        combined.add_page_break()
        combined.add_heading(label, level=1)
        combined.add_paragraph(f"Source file: {path}")
        for line in text_runs(path):
            if line:
                add_paragraph_copy(combined, line)
            else:
                combined.add_paragraph()
    out_path = OUT_DIR / "Vishnusahasranamam_PDF_OCR_Studio_corrected_vol_01_05.docx"
    combined.save(str(out_path))
    return out_path


def write_report(audits: list[VolumeAudit], combined_path: Path) -> Path:
    report = OUT_DIR / "Vishnusahasranamam_PDF_OCR_Studio_audit.md"
    lines: list[str] = []
    lines.append("# PDF_OCR_Studio Corrected Source Audit")
    lines.append("")
    lines.append("Source policy: old `Vishnusahasranamam_Dayananda.docx` was not used.")
    lines.append("")
    lines.append(f"Combined source candidate: `{combined_path}`")
    lines.append("")
    lines.append("## Summary")
    lines.append("")
    lines.append("| Volume | Non-empty paragraphs | Characters | Devanagari chars | Latin chars | Suspect lines |")
    lines.append("|---|---:|---:|---:|---:|---:|")
    for audit in audits:
        lines.append(
            f"| {audit.label} | {audit.nonempty:,} | {audit.chars:,} | "
            f"{audit.devanagari_chars:,} | {audit.latin_chars:,} | {len(audit.suspect_lines):,} |"
        )
    lines.append("")
    total_suspect = sum(len(a.suspect_lines) for a in audits)
    if total_suspect:
        lines.append("## Suspect Lines")
        lines.append("")
        lines.append("These are audit flags, not automatic corrections.")
        lines.append("")
        for audit in audits:
            lines.append(f"### {audit.label}")
            lines.append("")
            if not audit.suspect_lines:
                lines.append("No obvious OCR artifact patterns found.")
                lines.append("")
                continue
            for line_no, reason, text in audit.suspect_lines[:40]:
                safe_text = text.replace("|", "\\|")
                lines.append(f"- Paragraph {line_no}: {reason}: {safe_text}")
            if len(audit.suspect_lines) > 40:
                lines.append(f"- ... {len(audit.suspect_lines) - 40} more flagged lines")
            lines.append("")
    lines.append("## Opening Samples")
    lines.append("")
    for audit in audits:
        lines.append(f"### {audit.label}")
        lines.append("")
        for sample in audit.samples:
            lines.append(f"> {sample}")
        lines.append("")
    report.write_text("\n".join(lines), encoding="utf-8")
    return report


def main() -> None:
    missing = [str(path) for _, path in VOLUMES if not path.exists()]
    if missing:
        raise SystemExit("Missing expected corrected DOCX files:\n" + "\n".join(missing))
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    audits = [audit_volume(label, path) for label, path in VOLUMES]
    combined_path = make_combined_docx()
    report_path = write_report(audits, combined_path)
    print(combined_path)
    print(report_path)
    for audit in audits:
        print(
            f"{audit.label}: {audit.nonempty} nonempty paragraphs, "
            f"{audit.chars} chars, {len(audit.suspect_lines)} suspect lines"
        )


if __name__ == "__main__":
    main()
