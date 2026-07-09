from __future__ import annotations

import argparse
import re
import subprocess
import tempfile
from dataclasses import dataclass
from concurrent.futures import ProcessPoolExecutor, as_completed
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


MIXED_SCRIPT_RE = re.compile(
    r"(?:[\u0900-\u097F][\u0900-\u097F\u200c\u200d\u00b7|/\\:;,.!?'’`~@#$%^&*+=_-]*[A-Za-z]"
    r"|[A-Za-z][A-Za-z0-9@#$%^*_+=~`'’.-]*[\u0900-\u097F])"
)
OCR_SYMBOL_RE = re.compile(r"[@#$%^*_+=~`]")
BROKEN_DEV_RE = re.compile(r"(?:^|\s)[\u093e-\u094d\u0951-\u0957](?:\s|$)")
REPEATED_DEV_RE = re.compile(r"([\u0900-\u097F]{3,})(?:\1){4,}")
CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


@dataclass
class PageText:
    volume: int
    scan_page: int
    text: str


def volume_number(path: Path) -> int:
    match = re.search(r"vol\s*0?(\d+)", path.name, re.IGNORECASE)
    if not match:
        raise ValueError(f"Could not determine volume from {path.name}")
    return int(match.group(1))


def pdf_to_pages(path: Path) -> list[str]:
    output = subprocess.check_output(["pdftotext", "-layout", str(path), "-"], text=True, errors="ignore")
    pages = output.split("\f")
    return [page.rstrip() for page in pages if page.strip()]


def pdf_page_count(path: Path) -> int:
    output = subprocess.check_output(["pdfinfo", str(path)], text=True, errors="ignore")
    match = re.search(r"^Pages:\s*(\d+)", output, re.MULTILINE)
    if not match:
        raise ValueError(f"Could not determine page count for {path}")
    return int(match.group(1))


def tesseract_page(task: tuple[str, int, int, int, str]) -> PageText:
    pdf_path, volume, page, dpi, langs = task
    with tempfile.TemporaryDirectory(prefix="vishnu_scan_ocr_") as temp_dir:
        prefix = Path(temp_dir) / f"vol_{volume:02d}_page_{page:04d}"
        subprocess.run(
            [
                "pdftoppm",
                "-r",
                str(dpi),
                "-f",
                str(page),
                "-l",
                str(page),
                "-png",
                pdf_path,
                str(prefix),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        images = sorted(Path(temp_dir).glob("*.png"))
        if not images:
            raise RuntimeError(f"No image rendered for {pdf_path} page {page}")
        result = subprocess.run(
            [
                "tesseract",
                str(images[0]),
                "stdout",
                "-l",
                langs,
                "--psm",
                "6",
                "--oem",
                "1",
                "-c",
                "preserve_interword_spaces=1",
            ],
            check=True,
            text=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
    return PageText(volume=volume, scan_page=page, text=result.stdout.strip())


def fresh_ocr_pages(pdfs: list[Path], dpi: int, langs: str, workers: int) -> list[PageText]:
    tasks: list[tuple[str, int, int, int, str]] = []
    for pdf in pdfs:
        vol = volume_number(pdf)
        for page in range(1, pdf_page_count(pdf) + 1):
            tasks.append((str(pdf), vol, page, dpi, langs))

    pages: list[PageText] = []
    with ProcessPoolExecutor(max_workers=workers) as executor:
        futures = [executor.submit(tesseract_page, task) for task in tasks]
        for done, future in enumerate(as_completed(futures), start=1):
            page_text = future.result()
            pages.append(page_text)
            if done % 25 == 0 or done == len(futures):
                print(f"OCR pages complete: {done}/{len(futures)}", flush=True)
    return sorted(pages, key=lambda item: (item.volume, item.scan_page))


def clean_line(line: str) -> str:
    line = line.replace("\xa0", " ")
    line = CONTROL_RE.sub("", line)
    line = re.sub(r"[ \t]+", " ", line).strip()
    line = line.replace("।।", "।।")
    return line


def page_paragraphs(text: str) -> list[str]:
    lines = [clean_line(line) for line in text.splitlines()]
    paragraphs: list[str] = []
    current: list[str] = []
    for line in lines:
        if not line:
            if current:
                paragraphs.append("\n".join(current).strip())
                current = []
            continue
        current.append(line)
    if current:
        paragraphs.append("\n".join(current).strip())
    return [para for para in paragraphs if para]


def issue_labels(text: str) -> list[str]:
    labels: list[str] = []
    if MIXED_SCRIPT_RE.search(text):
        labels.append("mixed Devanagari/Latin token")
    if BROKEN_DEV_RE.search(text):
        labels.append("broken Devanagari sign")
    if REPEATED_DEV_RE.search(text):
        labels.append("repeated Devanagari sequence")
    if OCR_SYMBOL_RE.search(text):
        labels.append("OCR-like symbol")
    return labels


def set_normal_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal.font.size = Pt(11)
    for style_name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 11)):
        style = document.styles[style_name]
        style.font.name = "Arial Unicode MS"
        style.font.size = Pt(size)


def add_paragraph(document: Document, text: str) -> None:
    para = document.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.name = "Arial Unicode MS"
    run.font.size = Pt(11)


def build_docx(pages: list[PageText], out_path: Path, report_path: Path) -> None:
    document = Document()
    set_normal_style(document)
    document.add_heading("Vishnusahasranamam - Clean OCR Draft", level=1)
    document.add_paragraph("Generated from scanned PDF volumes only. Page headings preserve source volume and scan page.")

    report_lines = [
        "# Vishnusahasranamam Clean OCR Draft - Review Report",
        "",
        "This report lists lines that should be manually checked. The DOCX was generated from scanned PDFs only.",
        "",
    ]

    current_volume: int | None = None
    issue_count = 0
    for page in pages:
        if current_volume != page.volume:
            if current_volume is not None:
                document.add_page_break()
            current_volume = page.volume
            document.add_heading(f"Volume {page.volume:02d}", level=1)
        document.add_heading(f"Volume {page.volume:02d} - Scan page {page.scan_page}", level=2)
        for para_text in page_paragraphs(page.text):
            labels = issue_labels(para_text)
            if labels:
                issue_count += 1
                report_lines.append(
                    f"- Volume {page.volume:02d}, scan page {page.scan_page}: {', '.join(labels)}"
                )
                report_lines.append(f"  Text: {para_text[:500].replace(chr(10), ' / ')}")
            add_paragraph(document, para_text)

    report_lines.insert(3, f"Total review items: {issue_count}")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)
    report_path.write_text("\n".join(report_lines) + "\n", encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--scan-dir", required=True)
    parser.add_argument("--out-docx", required=True)
    parser.add_argument("--out-report", required=True)
    parser.add_argument("--source", choices=["embedded", "fresh-ocr"], default="embedded")
    parser.add_argument("--dpi", type=int, default=300)
    parser.add_argument("--langs", default="san+eng")
    parser.add_argument("--workers", type=int, default=4)
    args = parser.parse_args()

    scan_dir = Path(args.scan_dir)
    pdfs = sorted(scan_dir.glob("*.pdf"), key=volume_number)
    pages: list[PageText] = []
    if args.source == "fresh-ocr":
        pages = fresh_ocr_pages(pdfs, dpi=args.dpi, langs=args.langs, workers=args.workers)
    else:
        for pdf in pdfs:
            vol = volume_number(pdf)
            for page_index, text in enumerate(pdf_to_pages(pdf), start=1):
                pages.append(PageText(volume=vol, scan_page=page_index, text=text))
    build_docx(pages, Path(args.out_docx), Path(args.out_report))
    print(f"PDFs: {len(pdfs)}")
    print(f"Pages: {len(pages)}")
    print(f"DOCX: {args.out_docx}")
    print(f"Report: {args.out_report}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
